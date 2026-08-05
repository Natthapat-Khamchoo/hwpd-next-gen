"""
สร้างเอกสารจับกุมเป็นไฟล์ .docx ในเครื่อง โดยไม่ต้องต่อ Google Docs

## ทำไมไม่ใช้ Google Docs

ของเดิม (`docs_service`) คัดลอกแม่แบบบน Drive แล้วสั่ง Docs API แทนที่ข้อความ
ซึ่งแลกมาด้วยสี่อย่าง — ต้องต่อเน็ตออก Google, กินโควตา API ที่ทั้งระบบใช้ร่วมกันอยู่แล้ว,
ทิ้งสำเนาไว้บน Drive ทุกครั้งที่กดปุ่ม, และผูกกับบัญชี OAuth บัญชีเดียวที่ถ้าโดนเพิกถอน
ก็ใช้ไม่ได้ทั้งระบบ

ตัวนี้อ่านแม่แบบ .docx ที่เก็บไว้ในรีโป แทนที่ตัวยึด แล้วคืนไบต์ของไฟล์กลับไปเลย
ไม่แตะเครือข่ายเลยสักครั้ง

## เรื่องที่ต้องระวัง — Word ตัดข้อความข้าม run

Word เก็บข้อความในย่อหน้าเป็นชิ้น ๆ (`run`) และตัดตรงไหนก็ได้ตามประวัติการพิมพ์
`<<OFFENSE>>` จึงมักไม่ได้อยู่ใน run เดียว แต่กระจายเป็น `<<OFF` + `ENSE>>` การไล่
`run.text.replace()` ทีละ run จึงไม่เจออะไรเลย ทั้งที่เปิดไฟล์ดูด้วยตาเห็นตัวยึดชัด ๆ

`_replace_in_paragraph` จึงต่อข้อความทุก run เข้าด้วยกันก่อน หาตำแหน่งบนข้อความที่ต่อแล้ว
แล้วค่อยเขียนกลับเฉพาะ run ที่ทับช่วงนั้น run อื่นไม่ถูกแตะ รูปแบบตัวอักษรของส่วนที่เหลือ
ในย่อหน้าจึงอยู่ครบ
"""

import io
import logging
import os
import re
from typing import Any, Dict, Iterable, List, Tuple

logger = logging.getLogger(__name__)

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "templates")
MAIN_TEMPLATE = os.path.join(TEMPLATE_DIR, "arrest_record.docx")
M22_TEMPLATE = os.path.join(TEMPLATE_DIR, "arrest_m22_23.docx")

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

#: ตัวยึดที่ยังไม่ได้แทนที่จะถูกลบทิ้ง ไม่ใช่ปล่อยค้างไว้
#
# เอกสารที่มี `<<OFFENSE>>` โผล่กลางหน้าคือเอกสารที่เจ้าหน้าที่อาจหยิบไปพิมพ์ส่ง
# พนักงานสอบสวนโดยไม่ทันสังเกต ช่องว่างอ่านออกว่ายังไม่ได้กรอก แต่ตัวยึดอ่านไม่ออก
PLACEHOLDER_RE = re.compile(r"<<[^<>\n]{1,60}>>")


class TemplateMissing(RuntimeError):
    """ไม่พบไฟล์แม่แบบในรีโป"""


def is_available() -> bool:
    """ใช้งานได้ไหม — ต้องมีทั้งไลบรารีและไฟล์แม่แบบ"""
    try:
        import docx  # noqa: F401
    except ImportError:
        return False
    return os.path.exists(MAIN_TEMPLATE) and os.path.exists(M22_TEMPLATE)


def _spans(text: str, runs) -> List[Tuple[int, int, int]]:
    """ตำแหน่งเริ่ม-จบของแต่ละ run บนข้อความที่ต่อกันแล้ว คืน [(index, start, end)]"""
    out, cursor = [], 0
    for index, run in enumerate(runs):
        length = len(run.text)
        out.append((index, cursor, cursor + length))
        cursor += length
    return out


def _replace_in_paragraph(paragraph, values: Dict[str, str]) -> int:
    """
    แทนที่ตัวยึดทุกตัวในย่อหน้าเดียว คืนจำนวนที่แทนไป

    ทำงานบนข้อความที่ต่อจากทุก run แล้วเขียนกลับเฉพาะ run ที่ทับช่วงของตัวยึด
    เพื่อไม่ให้รูปแบบของส่วนอื่นในย่อหน้าหาย
    """
    runs = paragraph.runs
    if not runs:
        return 0

    joined = "".join(run.text for run in runs)
    if "<<" not in joined:
        return 0

    def resolve(match: re.Match) -> str:
        key = match.group(0)[2:-2]
        return str(values.get(key, ""))

    replaced = 0
    # ไล่จากท้ายมาหน้า ตำแหน่งของตัวยึดที่ยังไม่ได้แทนจึงไม่เลื่อนตาม
    for match in reversed(list(PLACEHOLDER_RE.finditer(joined))):
        start, end = match.span()
        new_text = resolve(match)
        replaced += 1

        remaining = new_text
        for index, run_start, run_end in _spans(joined, runs):
            if run_end <= start or run_start >= end:
                continue
            head = runs[index].text[: max(0, start - run_start)]
            tail = runs[index].text[max(0, end - run_start):]
            runs[index].text = head + remaining + tail
            remaining = ""  # ข้อความใหม่ทั้งก้อนไปอยู่ run แรกที่ทับ ที่เหลือแค่ตัดออก
        joined = "".join(run.text for run in runs)

    return replaced


def _paragraphs(document) -> Iterable:
    """ทุกย่อหน้าในเอกสาร รวมที่อยู่ในตาราง หัวกระดาษ และท้ายกระดาษ"""
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            yield from ncell.paragraphs
    for section in document.sections:
        for part in (section.header, section.footer):
            yield from part.paragraphs
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def render(template_path: str, values: Dict[str, Any]) -> bytes:
    """
    เปิดแม่แบบ แทนที่ตัวยึดทั้งหมด คืนไบต์ของไฟล์ .docx

    ค่าที่ไม่ได้ส่งมาจะกลายเป็นช่องว่าง ไม่ใช่ตัวยึดค้าง (ดูหมายเหตุที่ PLACEHOLDER_RE)
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise TemplateMissing("ยังไม่ได้ติดตั้ง python-docx (pip install -r requirements.txt)") from exc

    if not os.path.exists(template_path):
        raise TemplateMissing(f"ไม่พบไฟล์แม่แบบ {os.path.basename(template_path)} ในโฟลเดอร์ templates")

    clean = {k: ("" if v is None else str(v)) for k, v in values.items()}
    document = Document(template_path)

    total = 0
    for paragraph in _paragraphs(document):
        total += _replace_in_paragraph(paragraph, clean)

    buffer = io.BytesIO()
    document.save(buffer)
    logger.info("สร้าง %s แทนที่ตัวยึด %d จุด", os.path.basename(template_path), total)
    return buffer.getvalue()


def leftover_placeholders(data: bytes) -> List[str]:
    """
    ตัวยึดที่ยังค้างอยู่ในไฟล์ที่สร้างเสร็จแล้ว ใช้ในเทสและตอนตรวจ

    อ่านผ่าน python-docx ไม่ใช่ regex บน XML ดิบ — `<<OFFENSE>>` หน้าตาเหมือน XML tag
    การ strip tag ด้วย `<[^>]+>` จะกินตัวยึดไปด้วย แล้วรายงานว่า "ไม่มีตัวยึดค้าง" เสมอ
    ไม่ว่าไฟล์จะเป็นอย่างไร
    """
    from docx import Document

    document = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in _paragraphs(document))
    return sorted(set(PLACEHOLDER_RE.findall(text)))


# ------------------------------------------------- ประกอบเอกสารจับกุมทั้งชุด


def _officer_fields(form: Dict[str, Any]) -> Dict[str, str]:
    """
    ฟอร์มส่งชื่อฟิลด์สั้น (respOfficer) แต่แม่แบบอ่านชื่อยาว (respOfficerName)
    รับทั้งสองแบบ ไม่งั้นช่องเจ้าหน้าที่ในเอกสารจะว่างเปล่าโดยไม่มีอะไรฟ้อง
    """
    def pick(*names: str) -> str:
        for name in names:
            value = str(form.get(name, "")).strip()
            if value:
                return value
        return ""

    return {
        "RESPONSIBLE_OFFICER_NAME": pick("respOfficerName", "respOfficer"),
        "RESPONSIBLE_OFFICER_PHONE": pick("respOfficerPhone", "respPhone"),
        "NOTIFYING_OFFICER_NAME": pick("notifyOfficerName", "notifyOfficer"),
        "NOTIFYING_OFFICER_PHONE": pick("notifyOfficerPhone", "notifyPhone"),
    }


def _safe_name(text: str, fallback: str = "เอกสาร") -> str:
    cleaned = re.sub(r'[\\/:*?"<>|\x00-\x1f]', "_", str(text or "")).strip()
    return (cleaned or fallback)[:80]


def build_arrest_documents(
    form_data: Dict[str, Any],
    suspects: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    """
    สร้างเอกสารทั้งชุดเป็นไฟล์ .docx คืน [{'name', 'filename', 'data'(bytes), 'leftover'}]

    ได้เอกสารสองชนิดเหมือนของเดิม — บันทึกจับกุมหนึ่งฉบับรวมผู้ต้องหาทุกคน
    และบันทึก ม.22, 23 คนละฉบับต่อผู้ต้องหาหนึ่งคน

    `leftover` คือตัวยึดที่แม่แบบมีแต่โค้ดไม่ได้ส่งค่ามาให้ ปกติต้องเป็นลิสต์ว่าง
    ถ้าไม่ว่างแปลว่าแม่แบบกับโค้ดหลุดจากกัน ผู้เรียกควรเอาไปเตือน ไม่ใช่กลืนเงียบ
    """
    from app.services.report_service import format_thai_datetime

    record_date = format_thai_datetime(form_data.get("recordDate", ""))
    arrest_date = format_thai_datetime(form_data.get("arrestDate", ""))
    officers = _officer_fields(form_data)
    people = [s for s in (suspects or []) if isinstance(s, dict) and str(s.get("name", "")).strip()]

    day = str(form_data.get("arrestDate", "")).split("T")[0] or "ไม่ระบุวันที่"
    shared = {
        "RECORD_DATE": record_date,
        "ARREST_DATE": arrest_date,
        "ARREST_LOCATION": form_data.get("arrestLocation", ""),
        "DETENTION_LOCATION": form_data.get("detentionLocation", ""),
        "OFFENSE": form_data.get("offense", ""),
        **officers,
    }

    results: List[Dict[str, Any]] = []

    main_values = {
        **shared,
        "ALL_SUSPECTS": form_data.get("allSuspectsText", ""),
        "CIRCUMSTANCES": form_data.get("circumstances", ""),
    }
    data = render(MAIN_TEMPLATE, main_values)
    results.append({
        "name": "บันทึกจับกุม (รวมทุกคน)",
        "filename": f"บันทึกจับกุม_{_safe_name(day)}.docx",
        "data": data,
        "leftover": leftover_placeholders(data),
    })

    for person in people:
        name = str(person.get("name", "")).strip()
        values = {
            **shared,
            "SUSPECT_NAME": name,
            "SUSPECT_IDCARD": person.get("idCard", ""),
            "SUSPECT_NAT": person.get("nat", ""),
            "SUSPECT_AGE": person.get("age", ""),
            "SUSPECT_ADDRESS": person.get("address", ""),
            "SUSPECT_PHONE": person.get("phone", ""),
            "BRIEF_CIRCUM": form_data.get("briefCircumstances", "") or form_data.get("circumstances", ""),
        }
        data = render(M22_TEMPLATE, values)
        results.append({
            "name": f"ม.22,23 ของ {name}",
            "filename": f"บันทึก_ม22_23_{_safe_name(name)}.docx",
            "data": data,
            "leftover": leftover_placeholders(data),
        })

    logger.info("สร้างเอกสารจับกุมเป็น .docx แล้ว %d ฉบับ (ผู้ต้องหา %d คน)", len(results), len(people))
    return results
