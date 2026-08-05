"""
สร้างเอกสารจับกุมเป็น .docx ในเครื่อง ไม่ผ่าน Google Docs

จุดที่เทสหนักที่สุดคือ **การแทนที่ตัวยึดที่ถูก Word ตัดข้ามหลาย run** เพราะเป็นกับดัก
ที่ทำให้โค้ดดูเหมือนทำงานได้ตอนทดสอบด้วยไฟล์ที่สร้างจากสคริปต์ (ซึ่งมี run เดียว)
แล้วไปเงียบ ๆ กับแม่แบบจริงที่คนพิมพ์มาเอง — ตัวยึดยังค้างอยู่ในเอกสารที่ส่งให้
พนักงานสอบสวน โดยไม่มีอะไรฟ้อง
"""

import io
import os
import unittest

os.environ.setdefault("SESSION_SECRET", "test-secret-for-docx-tests")

from app.services import docx_service  # noqa: E402

HAS_DOCX = docx_service.is_available()


def make_doc(chunks_per_paragraph):
    """
    สร้างเอกสารทดสอบโดย **บังคับให้ข้อความถูกตัดเป็น run ตามที่กำหนด**

    เขียน `<<OFFENSE>>` ทีเดียวจะได้ run เดียวซึ่งไม่ตรงกับของจริง แม่แบบที่คนพิมพ์เอง
    มักถูกตัดกลางคำเพราะเคยแก้ตัวหนา เปลี่ยนฟอนต์ หรือพิมพ์ทับ
    """
    from docx import Document

    document = Document()
    for chunks in chunks_per_paragraph:
        paragraph = document.add_paragraph()
        for chunk in chunks:
            paragraph.add_run(chunk)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def text_of(data):
    from docx import Document

    return "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)


@unittest.skipUnless(HAS_DOCX, "ยังไม่ได้ติดตั้ง python-docx หรือไม่มีไฟล์แม่แบบ")
class TestPlaceholderReplacement(unittest.TestCase):
    def render_chunks(self, chunks_per_paragraph, values):
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            handle.write(make_doc(chunks_per_paragraph))
            path = handle.name
        try:
            return text_of(docx_service.render(path, values))
        finally:
            os.unlink(path)

    def test_ตัวยึดที่อยู่ใน_run_เดียวแทนที่ได้(self):
        out = self.render_chunks([["โดยกล่าวหาว่า <<OFFENSE>>"]], {"OFFENSE": "ลักทรัพย์"})
        self.assertIn("โดยกล่าวหาว่า ลักทรัพย์", out)

    def test_ตัวยึดที่ถูกตัดข้าม_run_ก็ต้องแทนที่ได้(self):
        # กรณีจริงของแม่แบบที่คนพิมพ์เอง — นี่คือเคสที่ run.text.replace() ทีละ run จะพลาด
        out = self.render_chunks([["โดยกล่าวหาว่า <<OFF", "ENSE>>"]], {"OFFENSE": "ลักทรัพย์"})
        self.assertIn("โดยกล่าวหาว่า ลักทรัพย์", out)
        self.assertNotIn("<<", out)

    def test_ตัวยึดที่ถูกตัดเป็นสามชิ้นก็ยังได้(self):
        out = self.render_chunks([["<<", "OFFEN", "SE>>"]], {"OFFENSE": "ลักทรัพย์"})
        self.assertEqual(out.strip(), "ลักทรัพย์")

    def test_หลายตัวยึดในย่อหน้าเดียว(self):
        out = self.render_chunks(
            [["วันที่ <<ARREST_DATE>> ที่ <<ARREST_LOCATION>> ข้อหา <<OFFENSE>>"]],
            {"ARREST_DATE": "5 ส.ค. 69", "ARREST_LOCATION": "ทล.1", "OFFENSE": "ลักทรัพย์"},
        )
        self.assertIn("วันที่ 5 ส.ค. 69 ที่ ทล.1 ข้อหา ลักทรัพย์", out)

    def test_ตัวยึดที่ไม่มีค่าส่งมากลายเป็นช่องว่าง_ไม่ค้างอยู่(self):
        # เอกสารที่มี <<OFFENSE>> โผล่กลางหน้าคือเอกสารที่อาจถูกพิมพ์ส่งไปโดยไม่ทันสังเกต
        out = self.render_chunks([["ข้อหา <<OFFENSE>> จบ"]], {})
        self.assertIn("ข้อหา  จบ", out)
        self.assertNotIn("<<", out)

    def test_ข้อความที่ไม่มีตัวยึดต้องไม่ถูกแตะ(self):
        original = "ผู้ต้องหาได้รับทราบสิทธิตามกฎหมาย ดังนี้"
        out = self.render_chunks([[original]], {"OFFENSE": "ลักทรัพย์"})
        self.assertIn(original, out)

    def test_รูปแบบตัวอักษรของ_run_อื่นในย่อหน้าไม่หาย(self):
        import tempfile

        from docx import Document

        document = Document()
        paragraph = document.add_paragraph()
        bold = paragraph.add_run("ข้อหา ")
        bold.bold = True
        paragraph.add_run("<<OFFENSE>>")
        buffer = io.BytesIO()
        document.save(buffer)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            handle.write(buffer.getvalue())
            path = handle.name
        try:
            data = docx_service.render(path, {"OFFENSE": "ลักทรัพย์"})
        finally:
            os.unlink(path)

        runs = Document(io.BytesIO(data)).paragraphs[0].runs
        self.assertTrue(runs[0].bold, "run ที่เป็นตัวหนาอยู่เดิมต้องยังหนาอยู่")


@unittest.skipUnless(HAS_DOCX, "ยังไม่ได้ติดตั้ง python-docx หรือไม่มีไฟล์แม่แบบ")
class TestLeftoverDetection(unittest.TestCase):
    def test_ตรวจตัวยึดค้างด้วย_python_docx_ไม่ใช่_regex_บน_xml_ดิบ(self):
        """
        `<<OFFENSE>>` หน้าตาเหมือน XML tag การ strip tag ด้วย `<[^>]+>` จะกินมันไปด้วย
        แล้วรายงานว่า "ไม่มีตัวยึดค้าง" เสมอไม่ว่าไฟล์จะเป็นอย่างไร — เคยพลาดมาแล้ว
        """
        import re

        data = make_doc([["ข้อหา <<OFFENSE>> จบ"]])
        self.assertEqual(docx_service.leftover_placeholders(data), ["<<OFFENSE>>"])

        # วิธีที่ผิด: strip tag แล้วค่อยหา — ต้องหาไม่เจอ ซึ่งเป็นเหตุผลที่ห้ามใช้
        import zipfile

        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
        stripped = re.sub(r"<[^>]+>", "", xml)
        self.assertEqual(re.findall(r"<<[^<>]{1,60}>>", stripped), [])


@unittest.skipUnless(HAS_DOCX, "ยังไม่ได้ติดตั้ง python-docx หรือไม่มีไฟล์แม่แบบ")
class TestArrestDocuments(unittest.TestCase):
    FORM = {
        "recordDate": "2026-08-05T14:30",
        "arrestDate": "2026-08-05T09:15",
        "arrestLocation": "ทางหลวงหมายเลข 1 กม.ที่ 570",
        "detentionLocation": "ส.ทล.1 กก.5 บก.ทล.",
        "offense": "ร่วมกันมียาเสพติดไว้ในครอบครองเพื่อจำหน่าย",
        "circumstances": "ตั้งจุดตรวจแล้วพบรถต้องสงสัย",
        "briefCircumstances": "ตรวจค้นพบยาเสพติดในรถ",
        "allSuspectsText": "นายทดสอบ หนึ่ง และ นายทดสอบ สอง",
        "respOfficer": "ด.ต. ทดสอบ ระบบ",
        "respPhone": "0810000001",
    }
    SUSPECTS = [
        {"name": "ทดสอบ หนึ่ง", "idCard": "1234567890123", "nat": "ไทย", "age": "32",
         "address": "99/1 ต.สามเงา", "phone": "0812345678"},
        {"name": "ทดสอบ สอง", "idCard": "9876543210987", "nat": "ไทย", "age": "28",
         "address": "12/3 ต.แม่สอด", "phone": "0898765432"},
    ]

    def test_ได้บันทึกจับกุมหนึ่งฉบับและ_ม22_ต่อผู้ต้องหาหนึ่งคน(self):
        docs = docx_service.build_arrest_documents(self.FORM, self.SUSPECTS)
        self.assertEqual(len(docs), 3)
        self.assertIn("บันทึกจับกุม", docs[0]["name"])
        self.assertIn("ทดสอบ หนึ่ง", docs[1]["name"])
        self.assertIn("ทดสอบ สอง", docs[2]["name"])

    def test_ไม่มีตัวยึดค้างในเอกสารสักฉบับ(self):
        for doc in docx_service.build_arrest_documents(self.FORM, self.SUSPECTS):
            with self.subTest(doc=doc["name"]):
                self.assertEqual(doc["leftover"], [], f"{doc['name']} ยังมีตัวยึดค้าง")

    def test_ค่าที่ส่งไปปรากฏในเอกสารจริง(self):
        docs = docx_service.build_arrest_documents(self.FORM, self.SUSPECTS)
        main = text_of(docs[0]["data"])
        self.assertIn("กม.ที่ 570", main)
        self.assertIn("ยาเสพติด", main)

        m22 = text_of(docs[1]["data"])
        self.assertIn("1234567890123", m22)
        self.assertIn("ทดสอบ หนึ่ง", m22)
        self.assertIn("0810000001", m22)

    def test_ผู้ต้องหาแต่ละคนได้เอกสารของตัวเอง_ไม่ปนกัน(self):
        docs = docx_service.build_arrest_documents(self.FORM, self.SUSPECTS)
        first, second = text_of(docs[1]["data"]), text_of(docs[2]["data"])
        self.assertIn("1234567890123", first)
        self.assertNotIn("9876543210987", first)
        self.assertIn("9876543210987", second)
        self.assertNotIn("1234567890123", second)

    def test_ไม่มีผู้ต้องหาก็ยังได้บันทึกจับกุมหนึ่งฉบับ(self):
        docs = docx_service.build_arrest_documents(self.FORM, [])
        self.assertEqual(len(docs), 1)

    def test_ชื่อไฟล์ตัดอักขระที่ใช้ไต่_path_ออก(self):
        docs = docx_service.build_arrest_documents(
            self.FORM, [{"name": "../../etc/passwd", "idCard": "1"}]
        )
        self.assertNotIn("/", docs[1]["filename"])
        self.assertNotIn("\\", docs[1]["filename"])

    def test_ไฟล์ที่ได้เป็น_docx_จริง(self):
        import zipfile

        for doc in docx_service.build_arrest_documents(self.FORM, self.SUSPECTS[:1]):
            with self.subTest(doc=doc["name"]):
                self.assertTrue(doc["filename"].endswith(".docx"))
                with zipfile.ZipFile(io.BytesIO(doc["data"])) as archive:
                    self.assertIn("word/document.xml", archive.namelist())


if __name__ == "__main__":
    unittest.main()
